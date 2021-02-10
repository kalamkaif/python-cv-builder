from docx import Document
from docx.shared import Inches
import pyttsx3

def speak(text):
    pyttsx3.speak(text)

speak('welcome, create your resume')

document = Document()

document.add_picture('C:\\Users\\kalamkaif\\Pictures\\Camera Roll\\me.jpg',
                     width=Inches(2.0), height=Inches(2.0))

speak('Enter your name : ')
name = input('Enter your name : ')
speak("Enter your mobile number : ")
ph_no = input("Enter your mobile number : ")
speak('enter your email id : ')
email = input('enter your email id : ')

document.add_paragraph(name + ' | ' + ph_no + ' | ' + email)
#About me
document.add_heading('About me')
speak('Tell about your self ? ')
about_me = input('Tell about your self ? ')
document.add_paragraph(about_me)
#experience
speak('Tell about ur work experience ')
document.add_heading('Work Experience')
p = document.add_paragraph()
speak('Enter Company name : ')
company = input('Enter Company : ')
from_date = input('From date : ')
to_date = input('To date : ')

p.add_run(company + ' ').bold = True
p.add_run(from_date + '-' + to_date + '\n').italic = True
experience_details = input('Describe your experience at : ' +company)
p.add_run(experience_details)

# more experience
while True:
    speak('Do you have more experience  ? yes or no : ')
    has_more_exp = input('Do you have more exp  ? yes or no : ')
    if has_more_exp.lower() == 'yes':
        p = document.add_paragraph()
        company = input('Enter Company : ')
        from_date = input('From date : ')
        to_date = input('To date : ')
        p.add_run(company + ' ').bold = True
        p.add_run(from_date + '-' + to_date + '\n').italic = True
        experience_details = input('Describe your experience at : ' + company)
        p.add_run(experience_details)

    else:
        break

#skills
speak('enter ur skills')
document.add_heading('Skills')
skills = input('Enter Skills')
p = document.add_paragraph(skills)
p.style = 'List Bullet'

while True:
    speak('do you want add more skills? yes or no ?')
    has_more_skills = input('do you want add more skills? yes or no ?')
    if has_more_skills == 'yes':
        skills = input('Enter Skills : ')
        p = document.add_paragraph(skills)
        p.style = 'List Bullet'
    else:
        break

#Footer
section = document.sections[0]
footer = section.footer
p = footer.paragraphs[0]
p.text = "CV generated using AmigosCode"
speak('your cv is completed')

document.save('cv.docx')
